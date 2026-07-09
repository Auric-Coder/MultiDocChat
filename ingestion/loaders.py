"""File loading, chunking, and attribution metadata helpers."""

from __future__ import annotations

import re
from io import BytesIO
from pathlib import Path
from typing import Iterable, List, Optional, Protocol

from langchain_community.document_loaders import Docx2txtLoader, PyPDFLoader, TextLoader
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter


CHUNK_SIZE = 900
CHUNK_OVERLAP = 125
SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}


class UploadedFile(Protocol):
    """Small Protocol matching the Streamlit UploadedFile surface we need."""

    name: str

    def getbuffer(self) -> memoryview:
        """Return uploaded file bytes."""


def _splitter() -> RecursiveCharacterTextSplitter:
    return RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        add_start_index=True,
    )


def _normalize_heading(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


def _markdown_headings(text: str) -> list[tuple[int, str]]:
    headings: list[tuple[int, str]] = []
    for match in re.finditer(r"(?m)^(#{1,6})\s+(.+?)\s*#*\s*$", text):
        headings.append((match.start(), _normalize_heading(match.group(2))))
    return headings


def _docx_headings(file_path: str) -> list[tuple[int, str]]:
    try:
        from docx import Document as DocxDocument
    except ImportError:
        return []

    headings: list[tuple[int, str]] = []
    cursor = 0
    doc = DocxDocument(file_path)

    for paragraph in doc.paragraphs:
        text = paragraph.text or ""
        style_name = paragraph.style.name if paragraph.style else ""
        if text.strip() and style_name.lower().startswith("heading"):
            headings.append((cursor, _normalize_heading(text)))
        cursor += len(text) + 1

    return headings


def _docx_headings_from_bytes(file_bytes: bytes) -> list[tuple[int, str]]:
    try:
        from docx import Document as DocxDocument
    except ImportError:
        return []

    headings: list[tuple[int, str]] = []
    cursor = 0
    doc = DocxDocument(BytesIO(file_bytes))

    for paragraph in doc.paragraphs:
        text = paragraph.text or ""
        style_name = paragraph.style.name if paragraph.style else ""
        if text.strip() and style_name.lower().startswith("heading"):
            headings.append((cursor, _normalize_heading(text)))
        cursor += len(text) + 1

    return headings


def _heading_for_start(headings: list[tuple[int, str]], start_index: int) -> Optional[str]:
    section_heading: Optional[str] = None
    for heading_start, heading in headings:
        if heading_start > start_index:
            break
        section_heading = heading
    return section_heading


def _apply_common_metadata(
    chunks: Iterable[Document],
    *,
    source_file: str,
    section_headings: Optional[list[tuple[int, str]]] = None,
) -> List[Document]:
    enriched_chunks: list[Document] = []

    for index, chunk in enumerate(chunks):
        metadata = dict(chunk.metadata)
        metadata["source"] = source_file
        metadata["source_file"] = source_file
        metadata["chunk_id"] = f"{Path(source_file).stem}-{index:04d}"

        if "page" in metadata:
            metadata["page_number"] = int(metadata["page"]) + 1
            metadata.pop("page", None)
        elif section_headings is not None:
            start_index = int(metadata.get("start_index", 0))
            metadata["section_heading"] = _heading_for_start(
                section_headings, start_index
            )
        else:
            metadata["section_heading"] = None

        chunk.metadata = metadata
        enriched_chunks.append(chunk)

    return enriched_chunks


def _load_pdf(file_path: str) -> List[Document]:
    docs = PyPDFLoader(file_path).load()
    chunks = _splitter().split_documents(docs)
    return _apply_common_metadata(chunks, source_file=Path(file_path).name)


def _load_docx(file_path: str) -> List[Document]:
    docs = Docx2txtLoader(file_path).load()
    chunks = _splitter().split_documents(docs)
    return _apply_common_metadata(
        chunks,
        source_file=Path(file_path).name,
        section_headings=_docx_headings(file_path),
    )


def _load_markdown(file_path: str) -> List[Document]:
    docs = TextLoader(file_path, encoding="utf-8").load()
    chunks = _splitter().split_documents(docs)
    return _apply_common_metadata(
        chunks,
        source_file=Path(file_path).name,
        section_headings=_markdown_headings(docs[0].page_content if docs else ""),
    )


def _load_text(file_path: str) -> List[Document]:
    docs = TextLoader(file_path, encoding="utf-8").load()
    chunks = _splitter().split_documents(docs)
    return _apply_common_metadata(chunks, source_file=Path(file_path).name)


def load_and_chunk(file_path: str) -> List[Document]:
    """Load a supported file and return attribution-ready chunks."""

    extension = Path(file_path).suffix.lower()

    if extension == ".pdf":
        return _load_pdf(file_path)
    if extension == ".docx":
        return _load_docx(file_path)
    if extension == ".md":
        return _load_markdown(file_path)
    if extension == ".txt":
        return _load_text(file_path)

    raise ValueError(
        f"Unsupported file type '{extension}'. Supported: "
        f"{', '.join(sorted(SUPPORTED_EXTENSIONS))}"
    )


def _load_pdf_bytes(file_bytes: bytes, source_file: str) -> List[Document]:
    from pypdf import PdfReader

    reader = PdfReader(BytesIO(file_bytes))
    docs: list[Document] = []

    for page_index, page in enumerate(reader.pages):
        docs.append(
            Document(
                page_content=page.extract_text() or "",
                metadata={"page": page_index},
            )
        )

    chunks = _splitter().split_documents(docs)
    return _apply_common_metadata(chunks, source_file=source_file)


def _load_docx_bytes(file_bytes: bytes, source_file: str) -> List[Document]:
    from docx import Document as DocxDocument

    docx = DocxDocument(BytesIO(file_bytes))
    text = "\n".join(paragraph.text for paragraph in docx.paragraphs)
    docs = [Document(page_content=text, metadata={})]
    chunks = _splitter().split_documents(docs)
    return _apply_common_metadata(
        chunks,
        source_file=source_file,
        section_headings=_docx_headings_from_bytes(file_bytes),
    )


def _load_text_bytes(file_bytes: bytes, source_file: str) -> List[Document]:
    text = file_bytes.decode("utf-8")
    docs = [Document(page_content=text, metadata={})]
    chunks = _splitter().split_documents(docs)
    return _apply_common_metadata(chunks, source_file=source_file)


def _load_markdown_bytes(file_bytes: bytes, source_file: str) -> List[Document]:
    text = file_bytes.decode("utf-8")
    docs = [Document(page_content=text, metadata={})]
    chunks = _splitter().split_documents(docs)
    return _apply_common_metadata(
        chunks,
        source_file=source_file,
        section_headings=_markdown_headings(text),
    )


def load_uploaded_file(uploaded_file: UploadedFile) -> List[Document]:
    """Load a Streamlit upload from bytes so no temporary path is needed later."""

    source_file = Path(uploaded_file.name).name
    extension = Path(source_file).suffix.lower()
    file_bytes = bytes(uploaded_file.getbuffer())

    if extension == ".pdf":
        return _load_pdf_bytes(file_bytes, source_file)
    if extension == ".docx":
        return _load_docx_bytes(file_bytes, source_file)
    if extension == ".md":
        return _load_markdown_bytes(file_bytes, source_file)
    if extension == ".txt":
        return _load_text_bytes(file_bytes, source_file)

    raise ValueError(
        f"Unsupported file type for '{uploaded_file.name}'. Supported: "
        f"{', '.join(sorted(SUPPORTED_EXTENSIONS))}"
    )


def process_uploaded_files(files: List[UploadedFile]) -> List[Document]:
    """Load and chunk Streamlit uploads without persisting temporary source paths."""

    all_chunks: list[Document] = []

    for uploaded_file in files:
        all_chunks.extend(load_uploaded_file(uploaded_file))

    return all_chunks


def print_chunk_preview(chunks: List[Document], max_chars: int = 220) -> None:
    """Print chunk metadata and content previews for pre-vector-store inspection."""

    for chunk in chunks:
        preview = " ".join(chunk.page_content.split())[:max_chars]
        print(f"\nchunk_id={chunk.metadata.get('chunk_id')}")
        print(f"metadata={chunk.metadata}")
        print(f"preview={preview}")


if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(
        description="Load sample files and print chunk attribution metadata."
    )
    parser.add_argument("files", nargs="+", help="Sample files to inspect.")
    args = parser.parse_args()

    sample_chunks: list[Document] = []
    for sample_file in args.files:
        sample_chunks.extend(load_and_chunk(sample_file))

    print_chunk_preview(sample_chunks)
